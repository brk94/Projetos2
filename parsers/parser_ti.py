# ======================================================================================
# Imports
# ======================================================================================

import io
import re
from datetime import datetime
from typing import Optional, List

# Libs de Scraper
from pdfminer.high_level import extract_text
import docx

# Nossos módulos locais
from projeto_api_sonae.models import ParsedReport, Milestone, KPI, AreaNegocioEnum
from projeto_api_sonae.utils import (
    helper_extrair_BLOCO_TEXTO,
    helper_extrair_VALOR_LINHA,
    helper_limpar_financeiro,
)
from .base import ReportParser

# -------- Utilidades  --------
_STATUS_REGEX = r"(Conclu[ií]do|Em\s+Andamento|Em\s+Risco|Atrasado|Planejado|Pendente)"

def _parse_data(s: str):
    try:
        return datetime.strptime(s.strip(), "%d/%m/%Y").date()
    except Exception:
        return None

def _formatar_data_string_ou_dash(s: str) -> Optional[str]:
    """
    Retorna:
      - '—' se vier vazio/traço,
      - 'dd/mm/aaaa' se a data for válida,
      - caso contrário, o texto original (limpo).
    """
    s = (s or "").strip()
    if not s or s in {"—", "-", "–", "‐", "―"}:
        return "—"
    d = _parse_data(s)
    return d.strftime("%d/%m/%Y") if d else s


# ======================================================================================
# PDF
# ======================================================================================
class IT_PDFReportParser(ReportParser):
    def parse(self, file_stream: io.BytesIO) -> Optional[ParsedReport]:
        # ------------------------------------------------------------------
        # 1) Extrai e normaliza texto do PDF
        # ------------------------------------------------------------------
        raw = extract_text(file_stream) or ""
        txt = (
            raw.replace("\r", "\n")
               .replace("\u00A0", " ")  # NBSP
               .replace("–", "-")
               .replace("—", "-")
        )

        # hifenização em fim de linha: "integra-\nção" -> "integração"
        txt = re.sub(r"-\n(?=\w)", "", txt)

        # cabeçalhos quebrados (comum em PDF)
        txt = txt.replace("Data\nRealizada", "Data Realizada")

        # espaços/quebras redundantes
        txt = re.sub(r"[ \t]+\n", "\n", txt)
        txt = re.sub(r"\n{3,}", "\n\n", txt)

        # ------------------------------------------------------------------
        # helpers internos
        # ------------------------------------------------------------------
        def _get_linha(pat: str, default: str = "") -> str:
            m = re.search(pat, txt, flags=re.IGNORECASE | re.MULTILINE)
            return m.group(1).strip() if m else default

        def _get_bloco(start_pat: str, stop_pat: str) -> str:
            h = re.search(start_pat, txt, flags=re.IGNORECASE | re.MULTILINE)
            if not h:
                return ""
            start = h.end()
            s = re.search(stop_pat, txt[start:], flags=re.IGNORECASE | re.MULTILINE)
            bloco = txt[start:(start + s.start())] if s else txt[start:]
            return bloco.strip()

        def normalizar_paragrafos_com_quebra_de_linha(text: str) -> str:
            """
            Junta 'quebras suaves' vindas do PDF.
            """
            if not text:
                return ""

            end_sentence = re.compile(r'[.!?;:…]["\')\]]*$')
            new_bullet   = re.compile(r"^\s*[•\-–]\s+")
            out, buf = [], ""

            for raw_ln in text.splitlines():
                ln = (raw_ln or "").strip()

                if not ln:
                    if buf:
                        out.append(buf.strip())
                        buf = ""
                    continue

                if new_bullet.match(ln):
                    if buf:
                        out.append(buf.strip())
                    buf = ln
                    continue

                if not buf:
                    buf = ln
                else:
                    if end_sentence.search(buf):
                        out.append(buf.strip())
                        buf = ln
                    else:
                        buf += " " + ln

            if buf:
                out.append(buf.strip())
            return "\n".join(out)

        # ------------------------------------------------------------------
        # 2) Cabeçalho / campos gerais
        # ------------------------------------------------------------------
        nome_projeto = _get_linha(r"Relat[oó]rio\s+de\s+Status\s+Semanal\s*[–—-]\s*([^\n]+)")
        codigo_projeto = _get_linha(r"(?:ID\s*do\s*Projeto|C[oó]digo)\s*:\s*([A-Z0-9\-\.]+)")
        gerente = _get_linha(r"Gerente\s+do\s+Projeto\s*:\s*([^\n]+?)(?:\s+Sprint\b|$)") \
                  or _get_linha(r"Reporte\s+de\s*:\s*([^\n]+)")
        sprint_str = _get_linha(r"Sprint\s*:\s*(?:Sprint\s*)?(\d+)", "0")
        try:
            numero_sprint = int(sprint_str)
        except Exception:
            numero_sprint = 0

        status_geral = _get_linha(r"Status\s+Geral(?:\s*\(.*?\))?\s*:\s*([^\n]+)", "N/A")
        if not status_geral or status_geral == "N/A":
            bloco_status = _get_bloco(
                r"(?im)^\s*\d+\.\s*Status\s+Geral(?:\s*\(.*?\))?\s*:?\s*$",
                r"^\s*\d+\.\s|^\Z"
            ) or ""
            linha = next((ln.strip() for ln in bloco_status.splitlines() if ln.strip()), "")
            if linha:
                status_geral = linha

        # ------------------------------------------------------------------
        # 3) Seções textuais
        # ------------------------------------------------------------------
        resumo_executivo = _get_bloco(r"^\s*\d+\.\s*Sum[áa]rio\s+Executivo\s*:?\s*$", r"^\s*\d+\.\s|^\Z")
        riscos_e_impedimentos = _get_bloco(r"^\s*\d+\.\s*Principais\s+Impedimentos\s+e\s+Riscos\s*:?\s*$", r"^\s*\d+\.\s|^\Z")
        proximos_passos = _get_bloco(r"^\s*\d+\.\s*Pr[oó]ximos\s+Passos\s*:?\s*$", r"^\s*\d+\.\s|^\Z")

        resumo_executivo      = normalizar_paragrafos_com_quebra_de_linha(resumo_executivo)
        riscos_e_impedimentos = normalizar_paragrafos_com_quebra_de_linha(riscos_e_impedimentos)
        proximos_passos       = normalizar_paragrafos_com_quebra_de_linha(proximos_passos)

        resumo_executivo      = re.sub(r"\n{2,}", "\n", resumo_executivo).strip()
        riscos_e_impedimentos = re.sub(r"\n{2,}", "\n", riscos_e_impedimentos).strip()
        proximos_passos       = re.sub(r"\n{2,}", "\n", proximos_passos).strip()

        # ------------------------------------------------------------------
        # 4) KPIs
        # ------------------------------------------------------------------
        orc_raw   = _get_linha(r"Or[çc]amento\s+Total(?:\s+do\s+Projeto)?\s*:\s*([^\n]+)")
        custo_raw = _get_linha(r"Custo\s+Realizado\s+(?:at[eé]\s+a\s+Data)?\s*:\s*([^\n]+)")
        kpis_data = [
            KPI(
                nome_kpi="Orçamento Total",
                categoria_kpi="Financeiro",
                valor_numerico_kpi=helper_limpar_financeiro(orc_raw),
                valor_texto_kpi=orc_raw or "",
            ),
            KPI(
                nome_kpi="Custo Realizado",
                categoria_kpi="Financeiro",
                valor_numerico_kpi=helper_limpar_financeiro(custo_raw),
                valor_texto_kpi=custo_raw or "",
            ),
        ]

        # ------------------------------------------------------------------
        # 5) Milestones — formato “parser-friendly” (PDF)
        # ------------------------------------------------------------------
        bloco_m = _get_bloco(r"^\s*\d+\.\s*Acompanhamento\s+de\s+Marcos\s*\(Milestones\)\s*:?\s*$",
                            r"^\s*\d+\.\s|^\Z")

        milestones = []
        if bloco_m:
            compact = re.sub(r"\n(?!Milestone:)", " ", bloco_m)
            compact = re.sub(r"[ \t]+", " ", compact).strip()

            for rec in re.findall(r"Milestone:\s.*?(?=(?:\s+Milestone:)|\Z)",
                                compact, flags=re.IGNORECASE | re.DOTALL):

                m = re.search(
                    rf"Milestone:\s*(?P<desc>.*?)\s*\|\s*Status:\s*(?P<status>{_STATUS_REGEX})"
                    r"\s*\|\s*Prevista:\s*(?P<prev>\d{2}/\d{2}/\d{4})\s*\|\s*Data\s+Realizada:\s*(?P<real>(?:\d{2}/\d{2}/\d{4}|—|–|-|‐|―))",
                    rec, flags=re.IGNORECASE | re.DOTALL
                )
                if not m:
                    continue

                desc        = re.sub(r"\s+", " ", m.group("desc")).strip()
                status      = re.sub(r"\s+", " ", m.group("status")).strip()
                dt_prev     = _parse_data(m.group("prev"))
                
                # CORREÇÃO: Converter data_real para date ou None
                dt_real_str = m.group("real").strip()
                if dt_real_str in ['—', '–', '-', '‐', '―']:
                    dt_real = None
                else:
                    dt_real = _parse_data(dt_real_str)

                milestones.append(Milestone(
                    descricao=desc,
                    status=status,
                    data_planejada=dt_prev,               # date
                    data_real_ou_revisada=dt_real         # date ou None
        ))

        return ParsedReport(
            codigo_projeto=codigo_projeto or "",
            nome_projeto=nome_projeto or "",
            gerente_projeto=gerente or "",
            numero_sprint=numero_sprint,
            status_geral=status_geral or "N/A",
            resumo_executivo=resumo_executivo or "",
            riscos_e_impedimentos=riscos_e_impedimentos or "",
            proximos_passos=proximos_passos or "",
            area_negocio=AreaNegocioEnum.TI,
            milestones=milestones,
            kpis=kpis_data,
            story_points_planejados=None,
            story_points_entregues=None,
        )

# ======================================================================================
# DOCX
# ======================================================================================
class IT_DocxReportParser(ReportParser):
    def parse(self, file_stream: io.BytesIO) -> Optional[ParsedReport]:
        doc = docx.Document(file_stream)

        # Junta parágrafos não vazios (mantido)
        full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])

        # ---------- Helpers locais (desta classe) ----------
        def _extrair_secao_por_paragrafos(doc_obj: docx.Document, header_pattern: str, stop_pattern: str) -> Optional[str]:
            header_re = re.compile(header_pattern, flags=re.IGNORECASE)
            stop_re   = re.compile(stop_pattern,  flags=re.IGNORECASE)
            capturing = False
            bucket = []
            for para in doc_obj.paragraphs:
                line = (para.text or "").strip()
                if not line:
                    if capturing:
                        bucket.append("")
                    continue
                if not capturing:
                    if header_re.search(line):
                        capturing = True
                    continue
                if stop_re.search(line):
                    break
                bucket.append(line)
            text = "\n".join(bucket).strip()
            return text or None

        def _normalizar_paragrafos(text: str) -> str:
            if not text:
                return ""
            end_sentence = re.compile(r'[.!?;:…]["\')\]]*$')
            new_bullet   = re.compile(r"^\s*[•\-–]\s+")
            out, buf = [], ""
            for raw_ln in text.splitlines():
                ln = (raw_ln or "").strip()
                if not ln:
                    if buf:
                        out.append(buf.strip()); buf = ""
                    continue
                if new_bullet.match(ln):
                    if buf: out.append(buf.strip())
                    buf = ln; continue
                if not buf: buf = ln
                else:
                    if end_sentence.search(buf):
                        out.append(buf.strip()); buf = ln
                    else:
                        buf += " " + ln
            if buf: out.append(buf.strip())
            return "\n".join(out)

        def _get_primeira_linha(text: str) -> str:
            for ln in (text or "").splitlines():
                ln = ln.strip()
                if ln: return ln
            return ""

        # CORREÇÃO: Função para parsear datas como objetos date
        def _parse_date_to_date(s: str) -> Optional[datetime.date]:
            """Converte string para objeto date, ou retorna None se inválida/vazia"""
            s = (s or "").strip()
            if not s or s in {"—", "–", "-", "‐", "―", ""}:
                return None
            
            # Normaliza separadores
            s = s.replace("-", "/")
            
            try:
                # Tenta parsear a data
                date_obj = datetime.strptime(s, "%d/%m/%Y")
                return date_obj.date()
            except ValueError:
                return None

        # -------------- Padrões --------------
        re_sumario_header   = r"^\s*(?:\d+\.\s*)?(?:sum[áa]rio\s+executivo|sum[áa]rio\s+de\s+atividades)\s*:?"
        re_riscos_header    = r"^\s*(?:\d+\.\s*)?(?:principais\s+impedimentos\s+e\s+riscos|impedimentos\s+e\s+riscos|riscos\s+e\s+impedimentos)\s*:?"
        re_proximos_header  = r"^\s*(?:\d+\.\s*)?(?:pr[oó]xim[oa]s?\s+passos|pr[oó]xim[oa]s?\s+a[çc][oõ]es)\s*:?"
        re_stop_any_header  = r"^\s*\d+\.\s|^\s*(?:sum[áa]rio|impedimentos|riscos|pr[oó]xim[oa]s)\b"

        # ---------- Dados genéricos ----------
        projeto_data = {
            "nome_projeto": helper_extrair_VALOR_LINHA(
                full_text, r"(?i)Relat[óo]rio de Status Semanal\s*[–—-]\s*(.*?)\n"
            ),
            "codigo_projeto": helper_extrair_VALOR_LINHA(full_text, r"(?i)ID do Projeto:\s*(PROJ-\d+)"),
            "gerente_projeto": helper_extrair_VALOR_LINHA(full_text, r"(?i)Gerente do Projeto:\s*(.*?)\s+Sprint:"),
            "numero_sprint": int(helper_extrair_VALOR_LINHA(full_text, r"(?i)Sprint:\s*(?:Sprint\s*)?(\d+)") or 0),
            "status_geral": helper_extrair_VALOR_LINHA(full_text, r"(?i)Status Geral\s*\(?.*?Sa[úu]de?.*?\)?:\s*(.*?)\n"),
            "resumo_executivo":
                _extrair_secao_por_paragrafos(doc, re_sumario_header, re_stop_any_header)
                or helper_extrair_BLOCO_TEXTO(
                    full_text,
                    r"(?ims)(?:^|\n)\s*(?:\d+\.\s*)?(?:sum[áa]rio\s+executivo|sum[áa]rio\s+de\s+atividades)\s*:?\s*([\s\S]*?)(?=^\s*\d+\.\s|^\s*(?:principais\s+impedimentos|impedimentos|riscos|pr[oó]xim[oa]s)\b|$)"
                ),
            "riscos_e_impedimentos":
                _extrair_secao_por_paragrafos(doc, re_riscos_header, re_stop_any_header)
                or helper_extrair_BLOCO_TEXTO(
                    full_text,
                    r"(?ims)(?:^|\n)\s*(?:\d+\.\s*)?(?:principais\s+impedimentos\s+e\s+riscos|impedimentos\s+e\s+riscos|riscos\s+e\s+impedimentos)\s*:?\s*([\s\S]*?)(?=^\s*\d+\.\s|^\s*(?:sum[áa]rio|pr[oó]xim[oa]s)\b|$)"
                ),
            "proximos_passos":
                _extrair_secao_por_paragrafos(doc, re_proximos_header, re_stop_any_header)
                or helper_extrair_BLOCO_TEXTO(
                    full_text,
                    r"(?ims)(?:^|\n)\s*(?:\d+\.\s*)?(?:pr[oó]xim[oa]s?\s+passos|pr[oó]xim[oa]s?\s+a[çc][oõ]es)\s*:?\s*([\s\S]*?)(?=^\s*\d+\.\s|^\s*(?:sum[áa]rio|impedimentos|riscos)\b|$)"
                ),
        }

        if not projeto_data.get("status_geral"):
            bloco_status = helper_extrair_BLOCO_TEXTO(
                full_text,
                r"(?ims)^\s*(?:\d+\.\s*)?Status\s+Geral(?:\s*\(.*?\))?\s*:?\s*$\s*([\s\S]*?)(?=^\s*\d+\.\s|$)"
            ) or ""
            projeto_data["status_geral"] = _get_primeira_linha(bloco_status) or "N/A"

        for k in ("resumo_executivo", "riscos_e_impedimentos", "proximos_passos"):
            projeto_data[k] = _normalizar_paragrafos(projeto_data.get(k) or "")
            projeto_data[k] = re.sub(r"\n{2,}", "\n", projeto_data[k]).strip()

        # ------------- KPIs -------------
        orcamento = helper_extrair_VALOR_LINHA(full_text, r"(?i)Or[çc]amento Total do Projeto:\s*(.*?)\n")
        custo     = helper_extrair_VALOR_LINHA(full_text, r"(?i)Custo Realizado at[eé] a Data:\s*(.*?)\n")

        kpis_data = [
            KPI(
                nome_kpi="Orçamento Total",
                categoria_kpi="Financeiro",
                valor_numerico_kpi=helper_limpar_financeiro(orcamento),
                valor_texto_kpi=orcamento,
            ),
            KPI(
                nome_kpi="Custo Realizado",
                categoria_kpi="Financeiro",
                valor_numerico_kpi=helper_limpar_financeiro(custo),
                valor_texto_kpi=custo,
            ),
        ]

        # ---------- Milestones ----------
        milestones_data: List[Milestone] = []

        # 1) Leitura por tabela
        try:
            tabela = doc.tables[0]
            for row in tabela.rows[1:]:
                try:
                    data_planned_date = _parse_date_to_date((row.cells[2].text or "").strip())
                except (ValueError, IndexError):
                    data_planned_date = None
                
                real_date = _parse_date_to_date((row.cells[3].text or "").strip())
                
                milestones_data.append(Milestone(
                    descricao=row.cells[0].text.strip(),
                    status=row.cells[1].text.strip(),
                    data_planejada=data_planned_date,   # date object
                    data_real_ou_revisada=real_date,    # date object
                ))
        except Exception:
            pass  # segue para os fallbacks textuais

        # 2) NOVA ABORDAGEM: Processamento direto por parágrafos
        print("DEBUG - Nova abordagem: processamento por parágrafos...")
        
        # Encontra o índice do parágrafo que contém "Acompanhamento de Marcos (Milestones)"
        start_index = -1
        for i, para in enumerate(doc.paragraphs):
            if "Acompanhamento de Marcos (Milestones)" in para.text:
                start_index = i + 1
                break
        
        if start_index != -1:
            print(f"DEBUG - Encontrada seção de milestones no parágrafo {start_index}")
            
            # Coleta todos os parágrafos seguintes até encontrar o próximo cabeçalho ou fim
            milestone_paras = []
            for i in range(start_index, len(doc.paragraphs)):
                text = doc.paragraphs[i].text.strip()
                if not text:
                    continue
                # Para se encontrar o próximo cabeçalho (número seguido de ponto)
                if re.match(r'^\d+\.', text):
                    break
                milestone_paras.append(text)
            
            print(f"DEBUG - {len(milestone_paras)} parágrafos de milestones encontrados")
            
            # Junta todos os parágrafos em um texto
            milestones_text = " ".join(milestone_paras)
            print(f"DEBUG - Texto completo das milestones: {repr(milestones_text)}")
            
            # Divide por ponto e vírgula para obter milestones individuais
            milestone_parts = milestones_text.split(';')
            print(f"DEBUG - {len(milestone_parts)} partes após split por ;")
            
            for i, part in enumerate(milestone_parts):
                part = part.strip()
                if not part:
                    continue
                    
                print(f"DEBUG - Processando parte {i}: {repr(part)}")
                
                # Regex simples para extrair os componentes
                match = re.search(
                    r'(.+?)\s*-\s*(Concluído|Em\s+Risco|Atrasado|Em\s+Andamento|Planejado|Pendente)\s*-\s*Prevista\s*:\s*(\d{1,2}-\d{1,2}-\d{4})\s*-\s*Data\s+Realizada\s*:\s*(\d{1,2}-\d{1,2}-\d{4}|)',
                    part
                )
                
                if match:
                    desc = match.group(1).strip()
                    status = match.group(2).strip()
                    prev_date = _parse_date_to_date(match.group(3))
                    real_date = _parse_date_to_date(match.group(4))
                    
                    print(f"DEBUG - ✅ Match parte {i}: '{desc}' | {status} | {prev_date} | {real_date}")
                    
                    milestones_data.append(Milestone(
                        descricao=desc,
                        status=status,
                        data_planejada=prev_date,
                        data_real_ou_revisada=real_date,
                    ))
                else:
                    print(f"DEBUG - ❌ Parte {i} não match: {part}")

        print(f"DEBUG - Total de milestones extraídas: {len(milestones_data)}")

        return ParsedReport(
            **projeto_data,
            area_negocio=AreaNegocioEnum.TI,
            milestones=milestones_data,
            kpis=kpis_data,
            story_points_planejados=None,
            story_points_entregues=None,
        )