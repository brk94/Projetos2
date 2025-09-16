from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks
from typing import List, Dict, Any
import sqlalchemy
from sqlalchemy import text
import io
import re  # Importamos o RegEx!
from datetime import datetime
import spacy

# --- BIBLIOTECAS DO SCRAPER ---
from pdfminer.high_level import extract_text
import docx  # Para ler .docx
import pandas as pd  # Para ler .xlsx

# --- BIBLIOTECAS E CONFIGURAÇÃO DA IA (NÍVEL 3 - GEMINI) ---
import google.generativeai as genai

# !!! COLE SUA API KEY AQUI !!!
GEMINI_API_KEY = "" 

try:
    genai.configure(api_key=GEMINI_API_KEY)
    print("Configuração da API Gemini (Nível 3) carregada com sucesso!")
except Exception as e:
    print(f"ERRO AO CONFIGURAR API GEMINI: {e} - A IA Nível 3 (geração de resumo) não funcionará.")
    genai = None

# --- CARREGAR MODELO DE NLP (NÍVEL 2) ---
try:
    nlp = spacy.load("pt_core_news_md")
    print("Modelo de NLP (SpaCy Nível 2) carregado com sucesso!")
except OSError:
    print("ERRO: Modelo 'pt_core_news_md' do SpaCy não encontrado.")
    print("Rode: python -m spacy download pt_core_news_md")
    nlp = None 

# --- Configuração da Conexão com o DB Local ---
DATABASE_URL = "mysql+mysqlconnector://root:-MySQL1596@localhost:3306/mc_sonae_db"

try:
    engine = sqlalchemy.create_engine(DATABASE_URL)
    with engine.connect() as conn:
        print("Conexão com o Banco de Dados MySQL local estabelecida com sucesso!")
except Exception as e:
    print(f"ERRO: Não foi possível conectar ao Banco de Dados MySQL: {e}")
    exit()


app = FastAPI(title="API Projeto MC Sonae")


# --- NOSSAS DUAS (2) FUNÇÕES HELPER DE REGEX (A SOLUÇÃO DO BUG) ---

def helper_extrair_VALOR_LINHA(texto: str, padrao: str) -> str | None:
    """
    Helper "SNIPER": Caça padrões de UMA LINHA. (SEM re.DOTALL).
    Usado para campos como Status, Custo, Orçamento.
    """
    match = re.search(padrao, texto, re.IGNORECASE) # <-- SEM re.DOTALL
    if match:
        return match.group(1).strip()
    return None

def helper_extrair_BLOCO_TEXTO(texto: str, padrao: str) -> str | None:
    """
    Helper "REDE DE PESCA": Caça blocos de texto de MÚLTIPLAS LINHAS. (COM re.DOTALL).
    Usado para Sumário Executivo e Riscos.
    """
    match = re.search(padrao, texto, re.IGNORECASE | re.DOTALL) # <-- COM re.DOTALL
    if match:
        return match.group(1).strip()
    return None

# --- FUNÇÃO HELPER DE LIMPEZA FINANCEIRA ---
def helper_limpar_financeiro(texto: str) -> float:
    """Função helper para limpar texto de dinheiro e converter para float."""
    try:
        # Remove "R$", "R\$", pontos de milhar, e substitui vírgula decimal por ponto
        texto_limpo = re.sub(r"[R$\s\.]", "", str(texto)).replace(",", ".")
        # Pega apenas os números (caso tenha texto como "(Estouro)" junto)
        match = re.search(r"(-?[\d\.]+)", texto_limpo) # Adicionado suporte a números negativos
        if match:
            return float(match.group(1))
    except Exception:
        pass # Ignora erros de conversão (ex: se o input for None)
    return 0.0


# --- LISTAS DE PALAVRAS-CHAVE (IA NÍVEL 2) ---
LEMMAS_DE_RISCO = {'bloquear', 'impedimento', 'estourar', 'falhar', 'complexidade', 'risco', 'pressão', 'extra'}
LEMMAS_DE_ATRASO = {'atrasar', 'atrasado'} 


# --- FUNÇÃO GERADORA DE RESUMO (IA NÍVEL 3 - CORRIGIDA) ---
def gerar_resumo_com_gemini(projeto_data: dict, milestones_data: list) -> str:
    """Usa a IA Nível 3 (Gemini) para gerar um resumo do zero, baseado nos dados estruturados."""
    
    # --- Formatação de Strings (pt-BR) ---
    budget_str = f"R$ {projeto_data.get('budget_total', 0):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    cost_str = f"R$ {projeto_data.get('cost_realized', 0):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    variance_str = projeto_data.get('variance_text', 'N/A')
    fin_narrativa_str = projeto_data.get('financial_narrative', 'Nenhum contexto financeiro fornecido.')

    # Formata milestones
    milestones_texto_limpo = ""
    if milestones_data:
        for item in milestones_data:
            desc = item.get('description', 'N/A')
            status = item.get('status', 'N/A')
            milestones_texto_limpo += f"- Marco: '{desc}', está com status: '{status}'.\n"
    else:
        milestones_texto_limpo = "Nenhum milestone detalhado foi extraído."

    prompt = f"""
    Você é um Diretor de Projetos (PMO) sênior. Escreva um resumo executivo curto (máximo 4 frases), 
    profissional e direto (em português do Brasil) com base NOS SEGUINTES DADOS BRUTOS de um projeto:

    - Nome do Projeto: {projeto_data.get('project_name', 'N/A')}
    - Status Calculado (pela IA Nível 2): {projeto_data.get('overall_status', 'N/A')}
    - Sprint Atual: {projeto_data.get('last_sprint_reported', 'N/A')}
    
    - Dados Financeiros (Formato: pt-BR):
      - Orçamento Total: {budget_str}
      - Custo Realizado: {cost_str}
      - Variação (texto do relatório): {variance_str}

    - CONTEXTO FINANCEIRO (IMPORTANTE - A EXPLICAÇÃO DA VARIAÇÃO):
      {fin_narrativa_str}

    - Lista de Milestones (Metas do Projeto):
      {milestones_texto_limpo}

    Por favor, gere um "Sumário Executivo" que destaque os pontos principais. 
    Use o CONTEXTO FINANCEIRO para explicar corretamente o orçamento.
    Seja direto. NÃO inclua "Riscos e Impedimentos".
    """
    
    try:
        if not genai: 
            raise Exception("API Key do Gemini não foi configurada.")
            
        model = genai.GenerativeModel('gemini-2.5-pro') # Usando o Modelo PRO
        generation_config = genai.types.GenerationConfig(
            temperature=0.0
        )
        
        response = model.generate_content(prompt, generation_config=generation_config)
        return response.text.replace("*", "").strip()
        
    except Exception as e:
        print(f"ERRO AO CHAMAR API GEMINI (NÍVEL 3): {e}")
        return projeto_data.get('executive_summary', 'Falha ao gerar resumo de IA. Resumo original não encontrado.')

# --- FUNÇÕES DE PARSER (SCRAPERS NÍVEL 1) ---

def parse_pdf_report(pdf_stream: io.BytesIO) -> Dict[str, Any]:
    """ SCRAPER CORRIGIDO PARA PDF. """
    texto_bruto = extract_text(pdf_stream)
    
    # --- Campos de Linha Única (usando \n) ---
    nome_projeto = helper_extrair_VALOR_LINHA(texto_bruto, r"Status Report:\s*(.*?)\n")
    orcamento = helper_extrair_VALOR_LINHA(texto_bruto, r"Orçamento Total:\s*(.*?)\n")
    custo = helper_extrair_VALOR_LINHA(texto_bruto, r"Custo Realizado:\s*(.*?)\n")
    variance_text_original = helper_extrair_VALOR_LINHA(texto_bruto, r"Variação:\s*(.*?)\n")

    # --- Blocos de Texto (usando Lookahead) ---
    fin_narrativa = helper_extrair_BLOCO_TEXTO(texto_bruto, r"Visão Financeira \(IMPORTANTE\):\s*([\s\S]*?)(?=Orçamento Total:)")
    summary_raw = helper_extrair_BLOCO_TEXTO(texto_bruto, r"Sumário de Atividades:\s*([\s\S]*?)(?=Riscos e Impedimentos|Visão Financeira)")
    risks_raw = helper_extrair_BLOCO_TEXTO(texto_bruto, r"Riscos e Impedimentos:\s*([\s\S]*?)(?=Visão Financeira|Marcos do Projeto)")

    projeto_data = {
        "project_code": helper_extrair_VALOR_LINHA(texto_bruto, r"Código:\s*(PROJ-\d+)"),
        "project_name": nome_projeto, 
        "project_manager": helper_extrair_VALOR_LINHA(texto_bruto, r"Reporte de:\s*(.*?)\s+Data:"),
        "last_sprint_reported": int(helper_extrair_VALOR_LINHA(texto_bruto, r"Sprint:\s*Sprint\s*(\d+)") or 0),
        "overall_status_humano": helper_extrair_VALOR_LINHA(texto_bruto, r"Status Geral:\s*(.*?)\n"),
        
        "executive_summary": summary_raw,
        "risks_and_impediments": risks_raw,
        "financial_narrative": fin_narrativa, 
        
        "budget_total": helper_limpar_financeiro(orcamento),
        "cost_realized": helper_limpar_financeiro(custo),
        "variance_text": variance_text_original,
        "next_steps": None 
    }
    
    bloco_tabela = helper_extrair_BLOCO_TEXTO(texto_bruto, r"Marcos do Projeto \(Tabela\):\n([\s\S]*?)(?=\Z)")
    
    milestones_data = []
    if bloco_tabela:
        padrao_linha_tabela = re.compile(r"(\(.*?\))\s+(.*?)\s+(Concluído|Em\s+Andamento|Pendente)\s+(\d{2}/\d{2}/\d{4})", re.IGNORECASE)
        for match in padrao_linha_tabela.finditer(bloco_tabela):
            data_crua_str = match.group(4)
            try:
                data_obj = datetime.strptime(data_crua_str, '%d/%m/%Y')
                data_formatada_sql = data_obj.strftime('%Y-%m-%d')
            except ValueError:
                data_formatada_sql = None

            milestones_data.append({
                "description": f"{match.group(1)} {match.group(2).replace('\n', ' ')}",
                "status": match.group(3).replace('\n', ''),
                "date_planned": data_formatada_sql,
                "date_actual_or_revised": None 
            })
    
    # --- CHAMADA IA NÍVEL 2 (SPACY) ---
    texto_para_analise_ia = ((projeto_data["executive_summary"] or "") + " " + (projeto_data["risks_and_impediments"] or "") + " " + (projeto_data["variance_text"] or "") + " " + (projeto_data.get("financial_narrative", "") or "")).lower()
    status_calculado_pela_ia = (projeto_data["overall_status_humano"] or "Em Dia").strip()

    if nlp and texto_para_analise_ia: 
        doc = nlp(texto_para_analise_ia)
        lemmas_do_texto = {token.lemma_ for token in doc if not token.is_stop and not token.is_punct}
        if not lemmas_do_texto.isdisjoint(LEMMAS_DE_RISCO):
            status_calculado_pela_ia = "Em Risco"
        if not lemmas_do_texto.isdisjoint(LEMMAS_DE_ATRASO):
            status_calculado_pela_ia = "Atrasado"
    projeto_data["overall_status"] = status_calculado_pela_ia
    
    # --- CHAMADA IA NÍVEL 3 (GEMINI) ---
    print(f"Chamando IA Nível 3 (Gemini) para o projeto {projeto_data.get('project_code')}...")
    resumo_gerado_pela_ia = gerar_resumo_com_gemini(projeto_data, milestones_data) 
    projeto_data["executive_summary"] = resumo_gerado_pela_ia
    
    return {"projeto": projeto_data, "milestones": milestones_data}


def parse_docx_report(doc_stream: io.BytesIO) -> Dict[str, Any]:
    """ 
    SCRAPER CORRIGIDO PARA DOCX. 
    AGORA USA OS HELPERS 'VALOR_LINHA' E 'BLOCO_TEXTO' CORRETAMENTE.
    """
    doc = docx.Document(doc_stream)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""]) 

    # --- CAMPOS FINANCEIROS (COM REGEX \n SIMPLES, POIS ESTÃO EM LINHAS SEPARADAS) ---
    # CORREÇÃO: Usa o helper VALOR_LINHA (Sniper)
    orcamento = helper_extrair_VALOR_LINHA(full_text, r"Orçamento Total do Projeto:\s*(.*?)\n")
    custo = helper_extrair_VALOR_LINHA(full_text, r"Custo Realizado até a Data:\s*(.*?)\n")
    variance_text_original = helper_extrair_VALOR_LINHA(full_text, r"Variação \(Burn Rate\):\s*(.*?)\n")

    projeto_data = {
        # CORREÇÃO: Usa o helper VALOR_LINHA
        "project_code": helper_extrair_VALOR_LINHA(full_text, r"ID do Projeto:\s*(PROJ-\d+)"),
        "project_name": helper_extrair_VALOR_LINHA(full_text, r"Relatório de Status Semanal - (.*?)\n"), 
        "project_manager": helper_extrair_VALOR_LINHA(full_text, r"Gerente do Projeto:\s*(.*?)\s+Sprint:"),
        "last_sprint_reported": int(helper_extrair_VALOR_LINHA(full_text, r"Sprint:\s*Sprint\s*(\d+)") or 0),
        "overall_status_humano": helper_extrair_VALOR_LINHA(full_text, r"Status Geral \(Saúde\):\s*(.*?)\n"),
        
        # CORREÇÃO: Usa o helper BLOCO_TEXTO
        "executive_summary": helper_extrair_BLOCO_TEXTO(full_text, r"Sumário Executivo:\s*([\s\S]*?)(?=3\. Principais Impedimentos e Riscos|$)"),
        "risks_and_impediments": helper_extrair_BLOCO_TEXTO(full_text, r"Principais Impedimentos e Riscos:\s*([\s\S]*?)(?=4\. Próximos Passos|$)"),
        "next_steps": helper_extrair_BLOCO_TEXTO(full_text, r"Próximos Passos:\s*([\s\S]*?)(?=5\. Acompanhamento Financeiro|$)"),
        
        "financial_narrative": None,
        
        "budget_total": helper_limpar_financeiro(orcamento),
        "cost_realized": helper_limpar_financeiro(custo), # <-- Chave em INGLÊS (Correto)
        "variance_text": variance_text_original,
    }
    
    # (O resto da função: Tabela de Milestones, IA Nível 2, IA Nível 3, já estão corretos)
    milestones_data = []
    try:
        tabela = doc.tables[0] 
        for row in tabela.rows[1:]: 
            data_crua_str = row.cells[2].text.strip()
            try:
                data_obj = datetime.strptime(data_crua_str, '%d/%m/%Y')
                data_formatada_sql = data_obj.strftime('%Y-%m-%d')
            except ValueError:
                data_formatada_sql = None
            
            data_revisada_crua_str = row.cells[3].text.strip()
            if data_revisada_crua_str.lower() in ["(pendente)", ""]:
                data_revisada_sql = None
            else:
                try:
                    data_rev_obj = datetime.strptime(data_revisada_crua_str, '%d/%m/%Y')
                    data_revisada_sql = data_rev_obj.strftime('%Y-%m-%d')
                except ValueError:
                    data_revisada_sql = data_revisada_crua_str 

            milestones_data.append({
                "description": row.cells[0].text.strip(),
                "status": row.cells[1].text.strip(),
                "date_planned": data_formatada_sql, 
                "date_actual_or_revised": data_revisada_sql 
            })
    except IndexError:
        print(f"Aviso: DOCX {projeto_data.get('project_code', 'N/A')} não continha tabela de milestones.")
        
    # --- CHAMADA IA NÍVEL 2 (SPACY) ---
    LEMMAS_DE_RISCO = {'bloquear', 'impedimento', 'estourar', 'falhar', 'complexidade', 'risco', 'pressão', 'extra'}
    LEMMAS_DE_ATRASO = {'atrasar', 'atrasado'} 
    texto_para_analise_ia = ((projeto_data["executive_summary"] or "") + " " + (projeto_data["risks_and_impediments"] or "") + " " + (projeto_data["variance_text"] or "")).lower()
    status_calculado_pela_ia = (projeto_data["overall_status_humano"] or "Em Dia").strip()
    if nlp and texto_para_analise_ia:
        doc = nlp(texto_para_analise_ia)
        lemmas_do_texto = {token.lemma_ for token in doc if not token.is_stop and not token.is_punct}
        if not lemmas_do_texto.isdisjoint(LEMMAS_DE_RISCO):
            status_calculado_pela_ia = "Em Risco"
        if not lemmas_do_texto.isdisjoint(LEMMAS_DE_ATRASO):
            status_calculado_pela_ia = "Atrasado"
    projeto_data["overall_status"] = status_calculado_pela_ia

    # --- CHAMADA IA NÍVEL 3 (GEMINI) ---
    print(f"Chamando IA Nível 3 (Gemini) para o projeto {projeto_data.get('project_code')}...")
    resumo_gerado_pela_ia = gerar_resumo_com_gemini(projeto_data, milestones_data)
    projeto_data["executive_summary"] = resumo_gerado_pela_ia
    
    return {"projeto": projeto_data, "milestones": milestones_data}

def parse_xlsx_report(xls_stream: io.BytesIO) -> Dict[str, Any]:
    """ SCRAPER PARA XLSX (AINDA SÓ NÍVEL 1 - PRECISA DE UPGRADE) """
    
    df_resumo = pd.read_excel(xls_stream, sheet_name="Resumo_Executivo", header=None)
    resumo_dict = df_resumo.set_index(0)[1].to_dict()

    projeto_data = {
        "project_code": resumo_dict.get("ID do Projeto:"),
        "project_name": resumo_dict.get("Nome do Projeto:"),
        "project_manager": resumo_dict.get("Gerente do Projeto:"),
        "last_sprint_reported": int(resumo_dict.get("Sprint:", 0)),
        "overall_status": resumo_dict.get("Status Geral (Saúde):"), # <-- NOTA: Este está pegando o status humano
        "executive_summary": resumo_dict.get("Resumo:"), # <-- NOTA: Este está pegando o resumo humano
        "budget_total": helper_limpar_financeiro(str(resumo_dict.get("Orçamento Total:"))),
        "cost_realized": helper_limpar_financeiro(str(resumo_dict.get("Custo Realizado:"))),
        "risks_and_impediments": None, 
        "next_steps": None,          
        "variance_text": None,
        "financial_narrative": None # <-- Chave que faltava
    }
    
    df_milestones = pd.read_excel(xls_stream, sheet_name="Cronograma_Milestones")
    milestones_data_raw = df_milestones.to_dict('records')
    milestones_data = []
    for row in milestones_data_raw:
        milestones_data.append({
            "description": row.get('Descrição do Marco (Milestone)'),
            "status": row.get('Status do Marco'),
            "date_planned": row.get('Data Prevista').strftime('%Y-%m-%d') if pd.notna(row.get('Data Prevista')) else None,
            "date_actual_or_revised": None
        })
        
    # TODO: Esta função XLSX também precisa chamar a IA Nível 2 e Nível 3
    # (mas vamos deixar isso para depois, já que estamos focados no DOCX/PDF)

    return {"projeto": projeto_data, "milestones": milestones_data}


# --- O ENDPOINT DA API (O "GARÇOM") ---

@app.post("/processar-relatorios/")
async def processar_relatorios(files: List[UploadFile] = File(...)):
    """ Endpoint de Upload (com a lógica de 3 tabelas e chamadas de parser) """
    
    nomes_processados = []
    for file in files:
        dados_extraidos = None
        file_bytes = await file.read()
        file_stream = io.BytesIO(file_bytes)
        
        try:
            if file.filename.endswith(".pdf"):
                dados_extraidos = parse_pdf_report(file_stream)
            elif file.filename.endswith(".docx"):
                dados_extraidos = parse_docx_report(file_stream)
            elif file.filename.endswith((".xlsx", ".xls")):
                dados_extraidos = parse_xlsx_report(file_stream)
            else:
                continue 

            if not dados_extraidos or not dados_extraidos["projeto"].get("project_code"):
                continue

            projeto_data = dados_extraidos["projeto"]
            milestones_data = dados_extraidos["milestones"]
            project_code = projeto_data["project_code"]

            with engine.begin() as conn: 
                
                # PASSO A: UPSERT Projeto "Pai"
                sql_upsert_projeto_pai = text("""
                    INSERT INTO Projetos (project_code, project_name, project_manager, budget_total)
                    VALUES (:project_code, :project_name, :project_manager, :budget_total)
                    ON DUPLICATE KEY UPDATE
                        project_name = VALUES(project_name),
                        project_manager = VALUES(project_manager),
                        budget_total = VALUES(budget_total);
                """)
                conn.execute(sql_upsert_projeto_pai, parameters=projeto_data)

                # PASSO B: INSERT Relatório Histórico
                sql_insert_relatorio = text("""
                    INSERT INTO Relatorios_Sprint (
                        project_code_fk, sprint_number, report_date, overall_status, 
                        executive_summary, risks_and_impediments, next_steps, 
                        cost_realized, variance_text, financial_narrative
                    ) VALUES (
                        :project_code, :last_sprint_reported, CURDATE(), :overall_status,
                        :executive_summary, :risks_and_impediments, :next_steps,
                        :cost_realized, :variance_text, :financial_narrative
                    );
                """)
                result = conn.execute(sql_insert_relatorio, parameters=projeto_data)

                # PASSO C: Pegar o Novo ID do Relatório
                new_report_id = result.lastrowid 

                # PASSO D: Inserir os Milestones Históricos
                if milestones_data:
                    sql_insert_milestones = text("""
                        INSERT INTO Milestones_Historico (
                            report_id_fk, description, status, date_planned, date_actual_or_revised
                        ) VALUES (:report_id_fk, :description, :status, :date_planned, :date_actual_or_revised)
                    """)
                    
                    milestones_para_salvar = []
                    for milestone in milestones_data:
                        milestones_para_salvar.append({
                            "report_id_fk": new_report_id,
                            "description": milestone.get("description"),
                            "status": milestone.get("status"),
                            "date_planned": milestone.get("date_planned"),
                            "date_actual_or_revised": milestone.get("date_actual_or_revised")
                        })
                    
                    conn.execute(sql_insert_milestones, parameters=milestones_para_salvar)
                
                print(f"SUCESSO: Histórico do Sprint {projeto_data['last_sprint_reported']} salvo para o Projeto {project_code} (Report ID: {new_report_id}).")
                nomes_processados.append(file.filename)

        except Exception as e:
            print(f"ERRO AO PROCESSAR O ARQUIVO {file.filename}: {e}")
            
    return {"status": "sucesso", "arquivos_processados": nomes_processados}

# --- ENDPOINTS GET (LEITURA) ---

@app.get("/dashboard-executivo/")
async def get_dashboard_data():
    """ Dashboard Executivo (COM A CORREÇÃO PT-BR -> EN) """

    # CORREÇÃO FINAL: Esta query AGORA usa 'cost_realized' (Inglês) para bater com o DB
    query = text("""
        WITH UltimoSprint AS (
            SELECT 
                project_code_fk,
                MAX(sprint_number) AS max_sprint
            FROM Relatorios_Sprint
            GROUP BY project_code_fk
        ),
        StatusAtual AS (
            SELECT 
                rs.*
            FROM Relatorios_Sprint rs
            JOIN UltimoSprint us 
              ON rs.project_code_fk = us.project_code_fk 
             AND rs.sprint_number = us.max_sprint
        )
        SELECT
            COUNT(*) AS total_projetos,
            SUM(CASE WHEN overall_status = 'Em Dia' THEN 1 ELSE 0 END) AS projetos_em_dia,
            SUM(CASE WHEN overall_status = 'Em Risco' THEN 1 ELSE 0 END) AS projetos_em_risco,
            SUM(CASE WHEN overall_status = 'Atrasado' THEN 1 ELSE 0 END) AS projetos_atrasados,
            SUM(cost_realized) AS investimento_total_executado 
        FROM StatusAtual;
    """) # <-- CORRIGIDO para 'cost_realized' (Inglês)

    try:
        with engine.connect() as conn:
            resultado = conn.execute(query).fetchone()

            if resultado and resultado.total_projetos > 0:
                dados_reais = {
                    "total_projetos": int(resultado.total_projetos or 0),
                    "projetos_em_dia": int(resultado.projetos_em_dia or 0),
                    "projetos_em_risco": int(resultado.projetos_em_risco or 0),
                    "projetos_atrasados": int(resultado.projetos_atrasados or 0),
                    "investimento_total_executado": float(resultado.investimento_total_executado or 0.0)
                }
                return dados_reais

    except Exception as e:
        print(f"ERRO AO BUSCAR DADOS DO DASHBOARD: {e}")
        raise HTTPException(status_code=500, detail="Não foi possível buscar dados do banco de dados.")

    # Fallback
    return {
        "total_projetos": 0, "projetos_em_dia": 0, "projetos_em_risco": 0,
        "projetos_atrasados": 0, "investimento_total_executado": 0.0
    }

@app.get("/projetos/lista/")
async def get_projetos_lista():
    """ Retorna a lista de projetos (estático) """
    query = text("SELECT project_code, project_name FROM Projetos ORDER BY project_name;")
    try:
        with engine.connect() as conn:
            resultado = conn.execute(query).fetchall()
            projetos = [{"code": row.project_code, "name": row.project_name} for row in resultado]
            return projetos
    except Exception as e:
        print(f"ERRO AO BUSCAR LISTA DE PROJETOS: {e}")
        raise HTTPException(status_code=500, detail="Erro ao buscar lista de projetos.")


@app.get("/projeto/{project_code}/lista-sprints/")
async def get_sprints_do_projeto(project_code: str):
    """ Retorna a lista de sprints para um projeto """
    query = text("""
        SELECT report_id, sprint_number 
        FROM Relatorios_Sprint 
        WHERE project_code_fk = :code 
        ORDER BY sprint_number DESC;
    """)
    try:
        with engine.connect() as conn:
            resultado = conn.execute(query, parameters={"code": project_code}).fetchall()
            sprints = [{"report_id": row.report_id, "sprint_number": row.sprint_number} for row in resultado]
            return sprints
    except Exception as e:
        print(f"ERRO AO BUSCAR LISTA DE SPRINTS: {e}")
        raise HTTPException(status_code=500, detail="Erro ao buscar lista de sprints.")

@app.get("/relatorio/detalhe/{report_id}")
async def get_detalhe_do_relatorio(report_id: int):
    """ Endpoint CORRETO E FINAL para o Delta KPI (AGORA COM CÁLCULO DE SLIPPAGE) """
    
    query_relatorio_atual = text("SELECT * FROM Relatorios_Sprint WHERE report_id = :id;")
    
    try:
        with engine.connect() as conn:
            
            dados_relatorio_atual = conn.execute(query_relatorio_atual, parameters={"id": report_id}).fetchone()
            if not dados_relatorio_atual:
                raise HTTPException(status_code=404, detail="Relatório (Sprint) não encontrado")

            json_relatorio_atual = dict(dados_relatorio_atual._mapping)
            
            current_sprint_num = json_relatorio_atual.get('sprint_number', 0)
            current_proj_code = json_relatorio_atual.get('project_code_fk')
            
            status_anterior = None 
            custo_atual = json_relatorio_atual.get('cost_realized', 0.0) or 0.0 # INGLÊS

            # <--- MODIFICAÇÃO 1: Preparar o mapa de milestones anteriores ---
            prev_milestones_map = {} 

            if current_sprint_num > 1: 
                prev_sprint_num = current_sprint_num - 1
                
                # <--- MODIFICAÇÃO 2: A query agora TAMBÉM busca o report_id anterior ---
                query_relatorio_anterior = text("""
                    SELECT report_id, overall_status, cost_realized FROM Relatorios_Sprint 
                    WHERE project_code_fk = :code AND sprint_number = :sprint_num
                    LIMIT 1;
                """) # INGLÊS
                
                dados_relatorio_anterior = conn.execute(query_relatorio_anterior, 
                                                        parameters={"code": current_proj_code, "sprint_num": prev_sprint_num}
                                                       ).fetchone()
                
                if dados_relatorio_anterior:
                    json_relatorio_anterior = dict(dados_relatorio_anterior._mapping)
                    custo_anterior = json_relatorio_anterior.get('cost_realized', 0.0) or 0.0 # INGLÊS
                    cost_delta = custo_atual - custo_anterior 
                    status_anterior = json_relatorio_anterior.get('overall_status')
                    
                    # <--- MODIFICAÇÃO 3: Buscar os milestones do relatório anterior ---
                    prev_report_id = json_relatorio_anterior.get('report_id')
                    if prev_report_id:
                        query_prev_milestones = text("SELECT * FROM Milestones_Historico WHERE report_id_fk = :id;")
                        dados_prev_milestones_raw = conn.execute(query_prev_milestones, parameters={"id": prev_report_id}).fetchall()
                        # Criamos um mapa de "descrição" -> "objeto milestone" para consulta rápida
                        prev_milestones_map = {dict(m._mapping)['description']: dict(m._mapping) for m in dados_prev_milestones_raw}
                
                else:
                    cost_delta = custo_atual 
            
            else:
                cost_delta = custo_atual 

            query_projeto_pai = text("SELECT * FROM Projetos WHERE project_code = :code;")
            dados_projeto_pai = conn.execute(query_projeto_pai, parameters={"code": current_proj_code}).fetchone()

            # --- Busca dos Milestones ATUAIS ---
            query_milestones = text("SELECT * FROM Milestones_Historico WHERE report_id_fk = :id;")
            dados_milestones_raw = conn.execute(query_milestones, parameters={"id": report_id}).fetchall()

            projeto_json_pai = dict(dados_projeto_pai._mapping)
            projeto_json_final = {**projeto_json_pai, **json_relatorio_atual}
            
            projeto_json_final['cost_delta'] = cost_delta
            projeto_json_final['status_anterior'] = status_anterior

            # <--- MODIFICAÇÃO 4: Processar milestones ATUAIS e calcular slippage ---
            milestones_json_processed = []
            for m_raw in dados_milestones_raw:
                curr_m = dict(m_raw._mapping)
                curr_m['slippage'] = False # Default
                
                # Tenta encontrar o milestone correspondente no sprint anterior
                prev_m = prev_milestones_map.get(curr_m['description'])
                
                if prev_m:
                    # Um milestone "efetivo" é a data revisada, ou a planejada se a revisada não existir
                    curr_effective_date_str = curr_m.get('date_actual_or_revised') or curr_m.get('date_planned')
                    prev_effective_date_str = prev_m.get('date_actual_or_revised') or prev_m.get('date_planned')
                    
                    # Compara as datas (somente se ambas existirem)
                    if curr_effective_date_str and prev_effective_date_str:
                        try:
                            # Converte datas (que estão como YYYY-MM-DD no DB)
                            curr_date = datetime.strptime(curr_effective_date_str, '%Y-%m-%d').date()
                            prev_date = datetime.strptime(prev_effective_date_str, '%Y-%m-%d').date()
                            
                            if curr_date > prev_date:
                                curr_m['slippage'] = True # A data foi "empurrada"!
                        
                        except (ValueError, TypeError):
                            pass # Ignora erros de parsing de datas (ex: se for 'N/A')

                milestones_json_processed.append(curr_m)
            
            # <--- MODIFICAÇÃO 5: Retorna a lista de milestones PROCESSADA ---
            return {"projeto": projeto_json_final, "milestones": milestones_json_processed}
            
    except Exception as e:
        print(f"ERRO AO BUSCAR DETALHE DO RELATÓRIO {report_id}: {e}")
        raise HTTPException(status_code=500, detail="Erro ao buscar detalhes do relatório.")
    
# --- ENDPOINT DO GRÁFICO (CORRIGIDO PARA INGLÊS) ---

@app.get("/projeto/{project_code}/historico-financeiro/")
async def get_historico_financeiro(project_code: str):
    """ Endpoint CORRIGIDO (INGLÊS) para o gráfico de burn rate """
    
    query = text("""
        SELECT 
            rs.sprint_number, 
            rs.cost_realized,
            p.budget_total
        FROM Relatorios_Sprint rs
        JOIN Projetos p ON rs.project_code_fk = p.project_code
        WHERE rs.project_code_fk = :code
        ORDER BY rs.sprint_number ASC; 
    """)
    
    try:
        with engine.connect() as conn:
            resultado = conn.execute(query, parameters={"code": project_code}).fetchall()
            
            historico = [
                {
                    "sprint_number": row.sprint_number,
                    "cost_realized": row.cost_realized, # INGLÊS
                    "budget_total": row.budget_total
                } for row in resultado
            ]
            return historico
    except Exception as e:
        print(f"ERRO AO BUSCAR HISTÓRICO FINANCEIRO: {e}")
        raise HTTPException(status_code=500, detail="Erro ao buscar histórico financeiro.")
